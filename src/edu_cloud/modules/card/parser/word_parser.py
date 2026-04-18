"""Word 答案模板：生成骨架 + 解析答案 + 文字量→权重。"""
from __future__ import annotations

import logging
import math
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# --- 常量（与 layout.py 一致）---
CHARS_PER_LINE = 40        # 栏宽约 1120px / 28px字号 ≈ 40 字/行
LINE_HEIGHT = 35           # px，与 layout.py LINE_HEIGHT_PX 一致
IMAGE_BASE_HEIGHT = 200    # px，一张图的默认预留高度
MIN_QUESTION_HEIGHT = 80   # px，与 layout.py MIN_HEIGHT_PER_SUB 一致

# 题号正则（按优先级排列）
_HASH_PATTERN = re.compile(r'^#(\d+)[.、]')
# 自由格式：行首数字 + 各种分隔符（括号、点、顿号、空格后跟汉字/括号）
_FREE_PATTERN = re.compile(r'^(\d+)\s*[（(（.、,:：]')
# 带"题"字：第N题 / N题
_TI_PATTERN = re.compile(r'^第?(\d+)题')
# 选择题范围：1-6 ABCD...
_RANGE_PATTERN = re.compile(r'^(\d+)\s*[-–—]\s*(\d+)\s+(.+)')
# 原有兜底
_BARE_PATTERN = re.compile(r'^(\d+)[.、]')

# 段落标题（非答案内容）——匹配后跳过，不追加到上一题
_SECTION_HEADERS = re.compile(
    r'^(单选题?|多选题?|不定项选?择?题?|填空题?|解答题?|计算题?|'
    r'简答题?|实验题?|综合题?|选做题?|必做题?|附加题?|'
    r'一、|二、|三、|四、|五、|六、|七、|八、|九、|十、|'
    r'第[一二三四五六七八九十]+[部大]?题|'
    r'[IVⅠⅡⅢⅣⅤⅥivx]+[.、\s])\s*$',
    re.IGNORECASE,
)


def generate_word_template(questions: list[dict], output_path: str) -> None:
    """生成 Word 答案模板骨架，教师下载后填写答案。

    Args:
        questions: [{"number": int, "question_type": str}, ...]
        output_path: 输出 .docx 路径
    """
    doc = Document()
    sorted_qs = sorted(questions, key=lambda q: q["number"])
    for q in sorted_qs:
        doc.add_paragraph(f"#{q['number']}.")
    doc.save(output_path)


def parse_word_answers(
    file_path: str | Path,
    *,
    expected_numbers: set[int] | None = None,
) -> list[dict]:
    """解析教师填写后的 Word 答案文件。

    Args:
        file_path: .docx 文件路径
        expected_numbers: 数据库已知题号集合（兜底匹配用）

    Returns:
        [{"number": int, "answer_text": str, "image_count": int}, ...]
    """
    logger.info("parse_word_answers: file=%s, expected=%s", file_path, expected_numbers)
    doc = Document(str(file_path))

    # 先收集所有段落文本
    paragraphs: list[tuple[str, int]] = []  # (text, image_count)
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            text = "".join(
                (node.text or "")
                for node in element.iter(qn("w:t"))
            )
            text = text.strip()
            images_in_para = len(element.findall(f".//{qn('w:drawing')}")) + \
                             len(element.findall(f".//{qn('w:pict')}"))
            if text or images_in_para:
                paragraphs.append((text, images_in_para))

    results = _match_paragraphs(paragraphs, expected_numbers)
    logger.info("parse_word_answers: found %d questions, numbers=%s",
                len(results), [r["number"] for r in results])
    return results



def _match_paragraphs(
    paragraphs: list[tuple[str, int]],
    expected_numbers: set[int] | None = None,
) -> list[dict]:
    """从段落列表中匹配题号和答案，docx 和 pdf 共用。"""
    results: list[dict] = []
    current: dict | None = None
    seen_numbers: set[int] = set()

    for text, images_in_para in paragraphs:
        if _SECTION_HEADERS.match(text):
            continue

        range_results = _try_expand_range(text, seen_numbers)
        if range_results:
            for r in range_results:
                seen_numbers.add(r["number"])
                results.append(r)
            current = results[-1]
            continue

        matched_num, answer_start = _try_match_question(
            text, seen_numbers, expected_numbers,
        )

        if matched_num is not None:
            seen_numbers.add(matched_num)
            answer = text[answer_start:].strip()
            current = {"number": matched_num, "answer_text": answer, "image_count": images_in_para}
            results.append(current)
        elif current is not None:
            if text:
                current["answer_text"] += "\n" + text
            current["image_count"] += images_in_para

    return results


def _try_expand_range(
    text: str, seen_numbers: set[int],
) -> list[dict] | None:
    """尝试拆分选择题范围行，如 '1-6 CACADC' → 6 道独立题。

    也处理多题号空格分隔行：'13 ABC    14 ABC    15 AD   16 BC'
    """
    # 多个范围在同一行（优先检查）：1-6 CACADC     7-12 BCBCCA
    multi_ranges = re.findall(r'(\d+)\s*[-–—]\s*(\d+)\s+([A-Za-z]+)', text)
    if len(multi_ranges) >= 2:
        results = []
        for start_s, end_s, answers_str in multi_ranges:
            start, end = int(start_s), int(end_s)
            expanded = _split_range_answers(start, end, answers_str)
            for r in expanded:
                if r["number"] not in seen_numbers:
                    results.append(r)
        if results:
            return results

    # 单范围格式：N-M 答案字符串
    m = _RANGE_PATTERN.match(text)
    if m:
        start, end = int(m.group(1)), int(m.group(2))
        answers_str = m.group(3).strip()
        if start < end <= start + 30:
            results = _split_range_answers(start, end, answers_str)
            if results:
                return results

    # 多题号行：13 ABC    14 ABC    15 AD   16 BC
    multi = re.findall(r'(\d+)\s+([A-Z]+)', text)
    if len(multi) >= 2:
        results = []
        for num_str, ans in multi:
            num = int(num_str)
            if num not in seen_numbers:
                results.append({"number": num, "answer_text": ans, "image_count": 0})
        if results:
            return results

    return None


def _split_range_answers(start: int, end: int, answers_str: str) -> list[dict]:
    """将 'CACADC' 拆分到 start..end 的各题。"""
    count = end - start + 1
    # 清理答案字符串（去空格）
    clean = re.sub(r'\s+', '', answers_str)
    # 每个字母是一个答案（单选）
    if len(clean) == count and clean.isalpha():
        return [
            {"number": start + i, "answer_text": clean[i], "image_count": 0}
            for i in range(count)
        ]
    # 答案可能用空格/逗号分隔
    parts = re.split(r'[,，\s]+', answers_str.strip())
    if len(parts) == count:
        return [
            {"number": start + i, "answer_text": parts[i], "image_count": 0}
            for i in range(count)
        ]
    # 无法拆分，整体作为一题
    return [{"number": start, "answer_text": answers_str, "image_count": 0}]


def _try_match_question(
    text: str,
    seen_numbers: set[int],
    expected_numbers: set[int] | None,
) -> tuple[int | None, int]:
    """尝试从文本中识别题号。返回 (题号, 答案文本起始位置) 或 (None, 0)。"""
    # 1. #N. 标准格式（最高优先级）
    m = _HASH_PATTERN.match(text)
    if m:
        return int(m.group(1)), m.end()

    # 2. 第N题 / N题
    m = _TI_PATTERN.match(text)
    if m:
        num = int(m.group(1))
        if num not in seen_numbers:
            return num, m.end()

    # 3. 自由格式：N（ / N、等（不含 N. 以避免步骤编号误匹配）
    m = _FREE_PATTERN.match(text)
    if m:
        num = int(m.group(1))
        sep_char = text[m.end() - 1] if m.end() > 0 else ''
        # N. 格式容易误匹配步骤编号，只有当前无已识别题时才允许
        is_dot = sep_char in '.．'
        if num not in seen_numbers and 1 <= num <= 50:
            if is_dot:
                # 如果已识别过更大的题号，说明这是答案内部编号
                max_seen = max(seen_numbers) if seen_numbers else 0
                if num < max_seen:
                    return None, 0
            return num, m.end()

    # 4. 兜底：裸数字 + 分隔符
    m = _BARE_PATTERN.match(text)
    if m:
        num = int(m.group(1))
        if expected_numbers:
            if num in expected_numbers and num not in seen_numbers:
                return num, m.end()
        elif 1 <= num <= 100 and num not in seen_numbers:
            return num, m.end()

    return None, 0


def compute_weights_from_text(questions: list[dict]) -> list[dict]:
    """文字量 → 归一化权重，直接喂给 allocate_by_weights。

    Args:
        questions: [{"number": int, "answer_text": str, "image_count": int}, ...]

    Returns:
        [{"number": int, "weight": float}, ...]
    """
    if not questions:
        return []

    raw_weights = []
    for q in questions:
        text_lines = math.ceil(max(len(q["answer_text"]), 1) / CHARS_PER_LINE)
        raw = text_lines * LINE_HEIGHT
        raw += q.get("image_count", 0) * IMAGE_BASE_HEIGHT
        raw = max(raw, MIN_QUESTION_HEIGHT)
        raw_weights.append({"number": q["number"], "raw": raw})

    total = sum(w["raw"] for w in raw_weights)
    result = [{"number": w["number"], "weight": w["raw"] / total} for w in raw_weights]
    logger.debug("compute_weights: %d questions, weights=%s",
                 len(result), [(w["number"], round(w["weight"], 3)) for w in result])
    return result
