"""Tests for Word answer template parser."""
import math
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches


def _make_word(tmp_path: Path, lines: list[str], images: dict[int, int] | None = None) -> Path:
    """Helper: create a Word doc with given lines and optional images.

    Args:
        lines: text lines to add as paragraphs
        images: {line_index: num_images} — after which line to insert images
    """
    doc = Document()
    for i, line in enumerate(lines):
        doc.add_paragraph(line)
        if images and i in images:
            for _ in range(images[i]):
                # Add a tiny 1x1 image
                from io import BytesIO
                from PIL import Image as PILImage
                buf = BytesIO()
                PILImage.new("RGB", (10, 10), "red").save(buf, "PNG")
                buf.seek(0)
                doc.add_picture(buf, width=Inches(1))
    path = tmp_path / "answers.docx"
    doc.save(str(path))
    return path


class TestParseWordAnswers:
    """Test parse_word_answers() — extract text+images per question."""

    def test_basic_extraction(self, tmp_path):
        """#N. prefix correctly splits questions."""
        from edu_cloud.modules.card.word_parser import parse_word_answers

        path = _make_word(tmp_path, [
            "#1. A",
            "#2. B",
            "#3. x=3",
        ])
        result = parse_word_answers(path)
        assert len(result) == 3
        assert result[0]["number"] == 1
        assert result[0]["answer_text"] == "A"
        assert result[1]["answer_text"] == "B"
        assert result[2]["answer_text"] == "x=3"

    def test_multiline_answer(self, tmp_path):
        """Continuation lines belong to current question."""
        from edu_cloud.modules.card.word_parser import parse_word_answers

        path = _make_word(tmp_path, [
            "#1. (1) 第一步",
            "     (2) 第二步",
            "#2. A",
        ])
        result = parse_word_answers(path)
        assert len(result) == 2
        assert "(1) 第一步" in result[0]["answer_text"]
        assert "(2) 第二步" in result[0]["answer_text"]

    def test_fallback_without_hash(self, tmp_path):
        """DB-known question numbers match even without # prefix."""
        from edu_cloud.modules.card.word_parser import parse_word_answers

        path = _make_word(tmp_path, [
            "#1. A",
            "2. B",       # no # prefix
            "#3. C",
        ])
        expected_numbers = {1, 2, 3}
        result = parse_word_answers(path, expected_numbers=expected_numbers)
        assert len(result) == 3
        assert result[1]["number"] == 2

    def test_no_false_positive_on_answer_numbering(self, tmp_path):
        """Internal numbering like '1. 先求x' should NOT start new question."""
        from edu_cloud.modules.card.word_parser import parse_word_answers

        path = _make_word(tmp_path, [
            "#13. (1) 解方程",
            "1. 先求x的值",
            "2. 代入得y",
            "#14. A",
        ])
        result = parse_word_answers(path)
        assert len(result) == 2
        assert "先求x" in result[0]["answer_text"]
        assert result[1]["number"] == 14

    def test_fallback_no_false_positive_when_already_seen(self, tmp_path):
        """Bare '1.' should NOT start new question if #1 was already matched."""
        from edu_cloud.modules.card.word_parser import parse_word_answers

        path = _make_word(tmp_path, [
            "#1. A",
            "#13. (1) 解方程",
            "1. 先求x的值",   # bare '1.' but #1 already seen → must be continuation
            "#14. A",
        ])
        expected_numbers = {1, 2, 13, 14}
        result = parse_word_answers(path, expected_numbers=expected_numbers)
        assert len(result) == 3  # #1, #13, #14
        assert "先求x" in result[1]["answer_text"]  # continuation of #13

    def test_image_counting(self, tmp_path):
        """Images are counted per question."""
        from edu_cloud.modules.card.word_parser import parse_word_answers

        # Create doc with image after #1
        path = _make_word(tmp_path, [
            "#1. 见图",
            "#2. A",
        ], images={0: 2})  # 2 images after line 0
        result = parse_word_answers(path)
        assert result[0]["image_count"] >= 1  # at least detected

    def test_empty_doc(self, tmp_path):
        """Empty document returns empty list."""
        from edu_cloud.modules.card.word_parser import parse_word_answers

        path = _make_word(tmp_path, [])
        result = parse_word_answers(path)
        assert result == []


class TestComputeWeights:
    """Test compute_weights_from_text() — text length to normalized weights."""

    def test_basic_normalization(self):
        from edu_cloud.modules.card.word_parser import compute_weights_from_text

        questions = [
            {"number": 1, "answer_text": "A", "image_count": 0},
            {"number": 2, "answer_text": "x" * 200, "image_count": 0},
        ]
        weights = compute_weights_from_text(questions)
        assert len(weights) == 2
        # Q2 has much more text, should get higher weight
        assert weights[1]["weight"] > weights[0]["weight"]
        # Weights sum to ~1.0
        total = sum(w["weight"] for w in weights)
        assert abs(total - 1.0) < 0.001

    def test_image_adds_weight(self):
        from edu_cloud.modules.card.word_parser import compute_weights_from_text

        q_no_img = {"number": 1, "answer_text": "A", "image_count": 0}
        q_with_img = {"number": 2, "answer_text": "A", "image_count": 2}
        weights = compute_weights_from_text([q_no_img, q_with_img])
        assert weights[1]["weight"] > weights[0]["weight"]

    def test_minimum_height(self):
        """Even short answers get minimum weight."""
        from edu_cloud.modules.card.word_parser import compute_weights_from_text

        questions = [
            {"number": i, "answer_text": "A", "image_count": 0}
            for i in range(1, 6)
        ]
        weights = compute_weights_from_text(questions)
        # All equal → each gets 0.2
        for w in weights:
            assert abs(w["weight"] - 0.2) < 0.01


class TestGenerateWordTemplate:
    """Test generate_word_template() — create Word skeleton for teachers."""

    def test_basic_generation(self, tmp_path):
        from edu_cloud.modules.card.word_parser import generate_word_template

        questions = [
            {"number": 1, "question_type": "choice"},
            {"number": 2, "question_type": "choice"},
            {"number": 3, "question_type": "essay"},
        ]
        path = tmp_path / "template.docx"
        generate_word_template(questions, str(path))
        assert path.exists()

        # Read back and verify content
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("#1." in t for t in texts)
        assert any("#2." in t for t in texts)
        assert any("#3." in t for t in texts)

    def test_question_ordering(self, tmp_path):
        from edu_cloud.modules.card.word_parser import generate_word_template

        questions = [
            {"number": 5, "question_type": "essay"},
            {"number": 1, "question_type": "choice"},
        ]
        path = tmp_path / "template.docx"
        generate_word_template(questions, str(path))
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Should be ordered by number
        idx_1 = next(i for i, t in enumerate(texts) if "#1." in t)
        idx_5 = next(i for i, t in enumerate(texts) if "#5." in t)
        assert idx_1 < idx_5
